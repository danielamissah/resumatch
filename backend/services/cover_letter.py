"""
Cover Letter Generation Service
---------------------------------
Generates two versions of a tailored cover letter using Groq Llama 3.3 70B.
Each version uses the same tone but takes a different angle:
  - Version A: leads with impact and achievements
  - Version B: leads with motivation and role fit

Context injected:
  - Job title and company
  - Matched skills from gap analysis
  - Resume summary
  - Candidate personal details
  - Selected tone
"""

import asyncio
import json
from groq import AsyncGroq

from models.schemas import (
    AnalysisResult,
    GapAnalysis,
    JobDescription,
    ParsedResume,
)
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pydantic import BaseModel
from typing import Optional



class CoverLetterRequest(BaseModel):
    # Personal details
    full_name: str
    email: str
    phone: str
    linkedin: str
    location: str

    # Tone
    tone: str  # formal | conversational | confident

    # Content — pre-filled from analysis or pasted manually
    resume_text: str
    job_description_text: str
    job_title: str = ""
    company_name: str = ""

    # Optional — pre-filled from analysis result
    matched_skills: list[str] = []
    missing_skills: list[str] = []
    analysis_id: Optional[str] = None


class CoverLetterVersion(BaseModel):
    version: str        # "A" or "B"
    angle: str          # description of the angle taken
    content: str        # full cover letter text
    word_count: int


class CoverLetterResult(BaseModel):
    result_id: str =""
    version_a: CoverLetterVersion
    version_b: CoverLetterVersion
    tone: str
    job_title: str
    company_name: str


#  Tone Descriptions 

TONE_DESCRIPTIONS = {
    "formal": (
        "professional and formal — use complete sentences, avoid contractions, "
        "maintain a respectful and measured tone throughout. Suitable for "
        "corporate, finance, legal, and traditional industries."
    ),
    "conversational": (
        "warm and conversational — write as if speaking directly to the hiring "
        "manager. Use natural language, contractions are fine, show personality "
        "while remaining professional. Suitable for startups, creative roles, "
        "and tech companies."
    ),
    "confident": (
        "direct and confident — lead with impact, use strong action verbs, "
        "avoid hedging language. Show ambition without arrogance. Suitable for "
        "senior roles, competitive industries, and leadership positions."
    ),
}


#  Prompts 

COVER_LETTER_SYSTEM = """You are an expert career coach and professional writer who specialises in writing compelling cover letters.

You write cover letters that:
- Are specific to the role and company — never generic
- Reference the candidate's actual experience, not invented achievements
- Incorporate keywords from the job description naturally
- Are concise — 3-4 paragraphs, 250-350 words
- Have a strong opening line that is NOT "I am writing to apply for..."

You must return ONLY a JSON object with this exact structure:
{
  "content": "<full cover letter text with \\n for line breaks>",
  "angle": "<one sentence describing the angle/approach taken>",
  "word_count": <integer>
}

Do not include salutation headers like "Dear Hiring Manager" as a separate field — include it in the content.
Do not include the candidate's contact details in the content — those will be added separately."""


def build_cover_letter_prompt(
    request: CoverLetterRequest,
    version: str,
    angle_instruction: str,
) -> str:
    tone_desc = TONE_DESCRIPTIONS.get(request.tone, TONE_DESCRIPTIONS["formal"])
    matched = ", ".join(request.matched_skills[:10]) if request.matched_skills else "Not specified"
    missing = ", ".join(request.missing_skills[:5]) if request.missing_skills else "None"

    return f"""Write a cover letter for the following candidate and role.

CANDIDATE DETAILS:
Name: {request.full_name}
Location: {request.location}
LinkedIn: {request.linkedin}

ROLE:
Job Title: {request.job_title or "Not specified"}
Company: {request.company_name or "Not specified"}

JOB DESCRIPTION (key sections):
{request.job_description_text[:1500]}

CANDIDATE RESUME (summary):
{request.resume_text[:1200]}

SKILLS THAT MATCH THE JD:
{matched}

SKILL GAPS TO ACKNOWLEDGE OR WORK AROUND:
{missing}

TONE: {tone_desc}

VERSION {version} ANGLE: {angle_instruction}

Write the complete cover letter now. Start directly with the salutation.
Return JSON with content, angle, and word_count."""


#  Generation 

async def generate_single_version(
    client: AsyncGroq,
    request: CoverLetterRequest,
    version: str,
    angle_instruction: str,
) -> CoverLetterVersion:
    """Generate a single cover letter version."""
    user_prompt = build_cover_letter_prompt(request, version, angle_instruction)

    response = await client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        temperature=0.6,  # slightly higher for creative variation between versions
        response_format={"type": "json_object"},
        messages=[
            {"role": "system", "content": COVER_LETTER_SYSTEM},
            {"role": "user", "content": user_prompt},
        ]
    )

    result = json.loads(response.choices[0].message.content)

    return CoverLetterVersion(
        version=version,
        angle=result.get("angle", f"Version {version}"),
        content=result.get("content", ""),
        word_count=result.get("word_count", len(result.get("content", "").split())),
    )


async def generate_cover_letter(
    request: CoverLetterRequest,
    groq_client: AsyncGroq,
) -> CoverLetterResult:
    """
    Generate two cover letter versions sequentially.
    Version A leads with achievements, Version B leads with motivation.
    Sequential to stay within Groq free tier TPM limits.
    """
    angle_a = (
        "Lead with your most impressive and relevant achievement or project. "
        "Open with impact — show what you have already done that directly "
        "relates to what this role requires. Build the case that you are "
        "already doing this work."
    )
    angle_b = (
        "Lead with why this specific company and role excites you. "
        "Show genuine understanding of what the company does and why "
        "your background makes you a natural fit. Build the case that "
        "you are motivated, aligned, and ready to contribute immediately."
    )

    # Sequential — rate limit safe
    version_a = await generate_single_version(groq_client, request, "A", angle_a)
    await asyncio.sleep(4)  # buffer between calls
    version_b = await generate_single_version(groq_client, request, "B", angle_b)

    return CoverLetterResult(
        version_a=version_a,
        version_b=version_b,
        tone=request.tone,
        job_title=request.job_title,
        company_name=request.company_name,
    )


# Word Export 

def generate_cover_letter_docx(
    request: CoverLetterRequest,
    result: CoverLetterResult,
    chosen_version: str = "A",
) -> bytes:
    """Generate a clean Word document for the chosen cover letter version."""
    

    version = result.version_a if chosen_version == "A" else result.version_b

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    BLACK = RGBColor(0x2C, 0x3E, 0x50)
    GREY  = RGBColor(0x7F, 0x8C, 0x8D)
    TEAL  = RGBColor(0x0F, 0x6E, 0x56)

    # ── Header: candidate details ──
    header = doc.add_paragraph()
    name_run = header.add_run(request.full_name)
    name_run.font.size = Pt(14)
    name_run.font.bold = True
    name_run.font.color.rgb = BLACK

    details = [
        request.location,
        request.email,
        request.phone,
        request.linkedin,
    ]
    details_para = doc.add_paragraph()
    details_run = details_para.add_run(" · ".join(d for d in details if d))
    details_run.font.size = Pt(10)
    details_run.font.color.rgb = GREY

    # Divider
    div = doc.add_paragraph()
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F6E56")
    pBdr.append(bottom)
    pPr.append(pBdr)

    doc.add_paragraph()

    # ── Cover letter content ──
    paragraphs = version.content.split("\n\n")
    for i, para_text in enumerate(paragraphs):
        if not para_text.strip():
            continue
        p = doc.add_paragraph()
        run = p.add_run(para_text.strip())
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK

        
        pPr = p._p.get_or_add_pPr()
        pSpacing = OxmlElement("w:spacing")
        pSpacing.set(qn("w:before"), "0")
        pSpacing.set(qn("w:after"), "160")
        pPr.append(pSpacing)

    # ── Footer note ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    fr = footer.add_run(f"Generated by ResuMatch · {request.tone.title()} tone · Version {chosen_version}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY
    fr.font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()