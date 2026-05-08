import io
import re
from typing import Union

import fitz
from docx import Document

from models.schemas import InputFormat, ParsedResume, ResumeSection

SECTION_PATTERNS = [
    r"^(experience|work experience|professional experience|employment history)",
    r"^(education|academic background|qualifications)",
    r"^(skills|technical skills|core competencies|technologies|tools)",
    r"^(projects|personal projects|academic projects|key projects)",
    r"^(publications|research|research experience|papers)",
    r"^(certifications|licenses|credentials)",
    r"^(summary|professional summary|profile|about me|objective)",
    r"^(awards|honors|achievements|accomplishments)",
    r"^(languages|spoken languages)",
    r"^(volunteering|community|extracurricular)",
]
COMPILED_PATTERNS = [re.compile(p, re.IGNORECASE | re.MULTILINE) for p in SECTION_PATTERNS]

SKILL_KEYWORDS = {
    "python","pytorch","tensorflow","scikit-learn","keras","huggingface",
    "langchain","langgraph","openai","transformers","mlflow","ray",
    "xgboost","lightgbm","catboost","numpy","pandas","matplotlib",
    "aws","azure","gcp","sagemaker","docker","kubernetes","terraform",
    "ci/cd","github actions","jenkins","airflow","kafka","redis",
    "fastapi","flask","django","node.js","express","postgresql",
    "mysql","mongodb","supabase","prisma","graphql","rest api",
    "next.js","react","typescript","javascript","tailwind","vue",
    "sql","spark","hadoop","dbt","looker","tableau","power bi",
    "computer vision","nlp","rag","llm","deep learning","machine learning",
    "reinforcement learning","generative ai","diffusion models",
}


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text_from_plain(text: str) -> str:
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\r', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def detect_section_heading(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    for pattern in COMPILED_PATTERNS:
        if pattern.match(line):
            return True
    if len(line) < 40 and line == line.upper() and len(line.split()) <= 4:
        return True
    return False


def extract_bullets(text: str) -> list[str]:
    bullets = []
    for line in text.split('\n'):
        line = line.strip()
        if line.startswith(('•', '-', '–', '▪', '*', '◦', '→')):
            cleaned = re.sub(r'^[•\-–▪*◦→]\s*', '', line).strip()
            if cleaned and len(cleaned) > 10:
                bullets.append(cleaned)
        elif re.match(r'^(Developed|Built|Designed|Implemented|Led|Managed|Created|Deployed|'
                      r'Researched|Published|Engineered|Optimised|Optimized|Trained|Fine-tuned|'
                      r'Architected|Delivered|Improved|Reduced|Increased|Automated)', line):
            if len(line) > 15:
                bullets.append(line)
    return bullets


def parse_sections(raw_text: str) -> list[ResumeSection]:
    lines = raw_text.split('\n')
    sections: list[ResumeSection] = []
    current_heading = "Summary"
    current_lines: list[str] = []

    for line in lines:
        if detect_section_heading(line):
            if current_lines:
                content = '\n'.join(current_lines).strip()
                sections.append(ResumeSection(title=current_heading, content=content, bullets=extract_bullets(content)))
            current_heading = line.strip().title()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        content = '\n'.join(current_lines).strip()
        sections.append(ResumeSection(title=current_heading, content=content, bullets=extract_bullets(content)))

    if not sections:
        sections = [ResumeSection(title="Full Resume", content=raw_text, bullets=extract_bullets(raw_text))]
    return sections


def extract_skills(raw_text: str) -> list[str]:
    text_lower = raw_text.lower()
    return sorted({skill for skill in SKILL_KEYWORDS if skill.lower() in text_lower})


def parse_resume(content: Union[bytes, str], format: InputFormat) -> ParsedResume:
    if format == InputFormat.PDF:
        raw_text = extract_text_from_pdf(content)
    elif format == InputFormat.DOCX:
        raw_text = extract_text_from_docx(content)
    elif format == InputFormat.TEXT:
        raw_text = extract_text_from_plain(content)
    else:
        raise ValueError(f"Unsupported format: {format}")

    if not raw_text.strip():
        raise ValueError("Could not extract any text from the provided resume")

    return ParsedResume(
        raw_text=raw_text,
        sections=parse_sections(raw_text),
        skills=extract_skills(raw_text),
        format=format,
        word_count=len(raw_text.split())
    )
