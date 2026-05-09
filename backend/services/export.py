"""
Word Redline Export Service
---------------------------
Generates a .docx file that visually mimics Word track changes:
- Original bullets: red + strikethrough
- Rewritten bullets: green + underline

This is a redline document, not true OOXML track changes.
It opens perfectly in Word, Google Docs, and LibreOffice
and is visually identical to what reviewers see in review mode.
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models.schemas import AnalysisResult


# ─── Colours ──────────────────────────────────────────────────────────────────

RED   = RGBColor(0xC0, 0x39, 0x2B)   # deleted text
GREEN = RGBColor(0x27, 0xAE, 0x60)   # inserted text
GREY  = RGBColor(0x7F, 0x8C, 0x8D)   # section headings
BLACK = RGBColor(0x2C, 0x3E, 0x50)   # normal text


def add_strikethrough(run):
    """Add strikethrough formatting to a run via direct XML."""
    rPr = run._r.get_or_add_rPr()
    strike = OxmlElement("w:strike")
    strike.set(qn("w:val"), "true")
    rPr.append(strike)


def add_paragraph_spacing(para, before: int = 0, after: int = 6):
    """Set paragraph spacing."""
    pPr = para._p.get_or_add_pPr()
    pSpacing = OxmlElement("w:spacing")
    pSpacing.set(qn("w:before"), str(before))
    pSpacing.set(qn("w:after"), str(after))
    pPr.append(pSpacing)


def set_left_indent(para, indent_twips: int = 360):
    """Set left indent on a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pInd = OxmlElement("w:ind")
    pInd.set(qn("w:left"), str(indent_twips))
    pPr.append(pInd)


# ─── Document Builder ─────────────────────────────────────────────────────────

def generate_redline_docx(result: AnalysisResult) -> bytes:
    """
    Build a Word redline document from an AnalysisResult.

    Structure:
      - Title + job info
      - ATS match score summary
      - Per-section redline bullets
      - Skill gaps appendix

    Returns raw .docx bytes ready to stream.
    """
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # ── Title ──
    title = doc.add_heading("Resume Redline Report", level=1)
    title.runs[0].font.color.rgb = BLACK
    title.runs[0].font.size = Pt(18)

    # ── Job info ──
    job_line = doc.add_paragraph()
    add_paragraph_spacing(job_line, after=2)
    r = job_line.add_run(f"Role: {result.job.title or 'Not specified'}")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

    if result.job.company:
        company_line = doc.add_paragraph()
        add_paragraph_spacing(company_line, after=2)
        r2 = company_line.add_run(f"Company: {result.job.company}")
        r2.font.size = Pt(11)
        r2.font.color.rgb = GREY

    # ── Match score summary ──
    doc.add_paragraph()
    score_para = doc.add_paragraph()
    add_paragraph_spacing(score_para, after=4)
    sr = score_para.add_run(f"Overall Match Score: {result.match_score.overall}/100 — {result.match_score.verdict}")
    sr.font.size = Pt(11)
    sr.font.bold = True
    sr.font.color.rgb = (
        RGBColor(0x27, 0xAE, 0x60) if result.match_score.overall >= 75
        else RGBColor(0xE6, 0x7E, 0x22) if result.match_score.overall >= 50
        else RGBColor(0xC0, 0x39, 0x2B)
    )

    summary_para = doc.add_paragraph()
    add_paragraph_spacing(summary_para, after=12)
    sumr = summary_para.add_run(result.match_score.summary)
    sumr.font.size = Pt(10)
    sumr.font.color.rgb = GREY
    sumr.font.italic = True

    # ── Legend ──
    legend_heading = doc.add_paragraph()
    add_paragraph_spacing(legend_heading, before=6, after=4)
    lh = legend_heading.add_run("Track Changes Legend")
    lh.font.bold = True
    lh.font.size = Pt(11)
    lh.font.color.rgb = BLACK

    legend_del = doc.add_paragraph()
    add_paragraph_spacing(legend_del, after=2)
    set_left_indent(legend_del, 360)
    ld = legend_del.add_run("Original bullet (removed)")
    ld.font.color.rgb = RED
    ld.font.size = Pt(10)
    add_strikethrough(ld)

    legend_ins = doc.add_paragraph()
    add_paragraph_spacing(legend_ins, after=12)
    set_left_indent(legend_ins, 360)
    li = legend_ins.add_run("Rewritten bullet (inserted)")
    li.font.color.rgb = GREEN
    li.font.size = Pt(10)
    li.font.underline = True

    # ── Redline bullets grouped by section ──
    rewrites = result.bullet_rewrites.rewrites
    by_section: dict[str, list] = {}
    for rw in rewrites:
        by_section.setdefault(rw.section, []).append(rw)

    for section_name, section_rewrites in by_section.items():
        # Section heading
        sec_para = doc.add_paragraph()
        add_paragraph_spacing(sec_para, before=16, after=6)
        sh = sec_para.add_run(section_name.upper())
        sh.font.bold = True
        sh.font.size = Pt(10)
        sh.font.color.rgb = GREY
        sh.font.all_caps = True

        # Horizontal rule effect via bottom border
        pPr = sec_para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

        for rw in section_rewrites:
            # Original — red strikethrough
            del_para = doc.add_paragraph(style="List Bullet")
            add_paragraph_spacing(del_para, after=2)
            del_run = del_para.add_run(rw.original)
            del_run.font.color.rgb = RED
            del_run.font.size = Pt(10)
            add_strikethrough(del_run)

            # Rewritten — green underline
            ins_para = doc.add_paragraph(style="List Bullet")
            add_paragraph_spacing(ins_para, after=8)
            ins_run = ins_para.add_run(rw.rewritten)
            ins_run.font.color.rgb = GREEN
            ins_run.font.size = Pt(10)
            ins_run.font.underline = True

            # Keywords added note
            if rw.keywords_added:
                kw_para = doc.add_paragraph()
                add_paragraph_spacing(kw_para, after=10)
                set_left_indent(kw_para, 360)
                kw_run = kw_para.add_run(f"Keywords added: {', '.join(rw.keywords_added)}")
                kw_run.font.size = Pt(9)
                kw_run.font.color.rgb = GREY
                kw_run.font.italic = True

    # ── Skill Gaps Appendix ──
    doc.add_page_break()
    appendix_title = doc.add_heading("Appendix: Skill Gap Analysis", level=2)
    appendix_title.runs[0].font.size = Pt(14)
    appendix_title.runs[0].font.color.rgb = BLACK

    for gap in result.gap_analysis.missing_skills:
        gap_para = doc.add_paragraph()
        add_paragraph_spacing(gap_para, before=10, after=2)

        # Skill name + priority
        name_run = gap_para.add_run(f"{gap.skill}  ")
        name_run.font.bold = True
        name_run.font.size = Pt(11)
        name_run.font.color.rgb = BLACK

        priority_color = RED if gap.priority == "critical" else RGBColor(0xE6, 0x7E, 0x22) if gap.priority == "important" else GREY
        pri_run = gap_para.add_run(f"[{gap.priority.upper()}]")
        pri_run.font.size = Pt(9)
        pri_run.font.color.rgb = priority_color
        pri_run.font.bold = True

        # Context
        ctx_para = doc.add_paragraph()
        add_paragraph_spacing(ctx_para, after=2)
        set_left_indent(ctx_para, 360)
        ctx_run = ctx_para.add_run(gap.context)
        ctx_run.font.size = Pt(10)
        ctx_run.font.color.rgb = GREY

        # Suggestion
        sug_para = doc.add_paragraph()
        add_paragraph_spacing(sug_para, after=2)
        set_left_indent(sug_para, 360)
        sug_run = sug_para.add_run(f"→ {gap.suggestion}")
        sug_run.font.size = Pt(10)
        sug_run.font.color.rgb = BLACK

        # YouTube link as plain text (Word doesn't render hyperlinks prettily in all versions)
        yt_para = doc.add_paragraph()
        add_paragraph_spacing(yt_para, after=12)
        set_left_indent(yt_para, 360)
        yt_text = f"Learn more: https://www.youtube.com/results?search_query={gap.skill.replace(' ', '+')}+tutorial"
        yt_run = yt_para.add_run(yt_text)
        yt_run.font.size = Pt(9)
        yt_run.font.color.rgb = RGBColor(0x27, 0x6F, 0xBF)
        yt_run.font.underline = True

    # ── Footer note ──
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    add_paragraph_spacing(footer_para, before=20)
    fr = footer_para.add_run("Generated by ResuMatch")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY
    fr.font.italic = True

    # ── Serialise to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()